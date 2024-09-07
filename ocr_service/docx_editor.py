import os

import re
from docx import Document
import json
import subprocess


def remove_non_ascii(string):
    return "".join(
        char
        for char in string
        if ord(char) < 128 or (ord(char) in [196, 214, 220, 228, 246, 252, 223])
    )




def parse_docx(docx_file):
    try:
        document = Document(docx_file)
    except:
        print("Problem with the file")
        return []
    unique_fields = []
    for p in document.paragraphs:
        subset = re.findall(r'\[(.*?)\]', p.text)
        unique_fields.extend(subset)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                subset = re.findall(r'\[(.*?)\]', cell.text)
                unique_fields.extend(subset)
    for section in document.sections:
        for p in section.header.paragraphs:
            subset = re.findall(r'\[(.*?)\]', p.text)
            unique_fields.extend(subset)
    for p in document.sections[0].first_page_header.paragraphs:
        subset = re.findall(r'\[(.*?)\]', p.text)
        unique_fields.extend(subset)
    unique_final = list(set(unique_fields))
    return unique_final


def edit_docx(docx_file, dict):
    try:
        document = Document(docx_file)
    except:
        print("Problem with the file")
        return []
    for p in document.paragraphs:
        for key, value in dict.items():
            if key in p.text:
                p.text = p.text.replace(f"[{key}]", value)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in dict.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(f"[{key}]", value)
    for section in document.sections:
        for p in section.header.paragraphs:
            for key, value in dict.items():
                if key in p.text:
                    p.text = p.text.replace(f"[{key}]", value)
    for p in document.sections[0].first_page_header.paragraphs:
        for key, value in dict.items():
            if key in p.text:                
                p.text = p.text.replace(f"[{key}]", value)
    return document


def apply_options(docx_file, dict):
    try:
        document = Document(docx_file)
    except:
        print("Problem with the file")
        return []
    for p in document.paragraphs:
        for key, value in dict.items():
            key_mod = "{" + key + "}"
            if key in p.text:
                p.text = p.text.replace(key_mod, value)
    return document