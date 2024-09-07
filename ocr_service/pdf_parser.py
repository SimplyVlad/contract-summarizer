import os
import pypdfium2 as pdfium
from nltk.tokenize import sent_tokenize
import re
from docx import Document
import json





def remove_non_ascii(string):
    return "".join(
        char
        for char in string
        if ord(char) < 128 or (ord(char) in [196, 214, 220, 228, 246, 252, 223])
    )




def parse_pdf_language(file_path, n_sent, language="german", token_limit=250):
    file_splits = []
    file_splits_vanilla = []
    try:
        pdf = pdfium.PdfDocument(file_path)
    except:
        print("Problem with the file")
        return []
    for page_num, page in enumerate(pdf):
        textpage = page.get_textpage()
        text_all = textpage.get_text_range()
        text_all = remove_non_ascii(text_all)
        text_sent_vanilla = sent_tokenize(text_all)
        text_sent = [
            str(sent).replace("\r\n", " ")
            for sent in text_sent_vanilla]
        all_text = " ".join(text_sent)
        all_text_vanilla = " ".join(text_sent_vanilla)
        file_splits.append(all_text)
        file_splits_vanilla.append(all_text_vanilla)
    full_text = " ".join(file_splits)
    full_text_vanilla = " ".join(file_splits_vanilla)
    return full_text
    



def parse_docx(docx_file):
    try:
        document = Document(docx_file)
    except:
        print("Problem with the file")
        return []
    unique_fields = []
    for p in document.paragraphs:
        subset = re.findall(r'\[(.*?)\]', p.text)
        unique_fields.extend(subset)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                subset = re.findall(r'\[(.*?)\]', cell.text)
                unique_fields.extend(subset)
    for section in document.sections:
        for p in section.header.paragraphs:
            subset = re.findall(r'\[(.*?)\]', p.text)
            unique_fields.extend(subset)
    for p in document.sections[0].first_page_header.paragraphs:
        subset = re.findall(r'\[(.*?)\]', p.text)
        unique_fields.extend(subset)
    unique_final = list(set(unique_fields))
    return unique_final


def edit_docx(docx_file, dict):
    try:
        document = Document(docx_file)
    except:
        print("Problem with the file")
        return []
    for p in document.paragraphs:
        for key, value in dict.items():
            if key in p.text:
                p.text = p.text.replace(f"[{key}]", value)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in dict.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(f"[{key}]", value)
    for section in document.sections:
        for p in section.header.paragraphs:
            for key, value in dict.items():
                if key in p.text:
                    p.text = p.text.replace(f"[{key}]", value)
    for p in document.sections[0].first_page_header.paragraphs:
        for key, value in dict.items():
            if key in p.text:                
                p.text = p.text.replace(f"[{key}]", value)
    return document
